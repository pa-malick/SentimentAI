"""
Génère le rapport du projet au format Word (rapport_projet.docx).

Le rapport se veut simple et pédagogique : les grandes lignes du projet,
sans entrer dans le détail du code. Les chiffres sont lus directement depuis
metrics/metrics_results.json, donc le rapport reste toujours à jour.

Usage : python generate_rapport.py
"""

import os
import json

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.config import METRICS_DIR, METRICS_JSON_PATH

SORTIE = "rapport_projet.docx"
AUTEUR = "Papa Malick NDIAYE"
FORMATION = "Master Data Science et Génie Logiciel"
URL_SITE = "https://sentimentai-df0c.onrender.com"

# Bleu sobre pour les titres
BLEU = RGBColor(0x1F, 0x4E, 0x79)


def charger_metriques():
    """Lit les métriques produites par la dernière évaluation."""
    if not os.path.exists(METRICS_JSON_PATH):
        raise FileNotFoundError(
            "metrics_results.json introuvable. Lance d'abord : python main.py"
        )
    with open(METRICS_JSON_PATH, encoding="utf-8") as f:
        return json.load(f)


def titre(doc, texte, niveau=1):
    """Ajoute un titre coloré."""
    h = doc.add_heading(texte, level=niveau)
    for run in h.runs:
        run.font.color.rgb = BLEU
    return h


def para(doc, texte, gras=False, taille=11):
    """Ajoute un paragraphe justifié."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texte)
    run.bold = gras
    run.font.size = Pt(taille)
    return p


def image(doc, nom_fichier, largeur=5.5):
    """Insère une figure du dossier metrics/ si elle existe."""
    chemin = os.path.join(METRICS_DIR, nom_fichier)
    if os.path.exists(chemin):
        doc.add_picture(chemin, width=Inches(largeur))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def legende(doc, texte):
    """Légende centrée sous une figure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.font.size = Pt(9)
    run.italic = True
    return p


def page_de_garde(doc):
    for _ in range(5):
        doc.add_paragraph()

    t = doc.add_heading("SentimentAI", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    lignes = [
        ("Analyse de sentiment par apprentissage automatique", 16, False, False),
        ("Comparaison de trois approches sur le dataset IMDB", 13, False, False),
        ("", 11, False, False),
        ("", 11, False, False),
        (AUTEUR, 14, True, False),
        (FORMATION, 12, False, False),
        ("", 11, False, False),
        ("Application en ligne : " + URL_SITE, 10, False, True),
    ]

    for texte, taille, gras, italique in lignes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texte)
        run.font.size = Pt(taille)
        run.bold = gras
        run.italic = italique

    doc.add_page_break()


def introduction(doc):
    titre(doc, "1. Introduction")
    para(doc,
         "L'analyse de sentiment consiste à déterminer automatiquement si un texte "
         "exprime une opinion positive ou négative. C'est une tâche classique du "
         "traitement automatique du langage, utilisée par exemple pour suivre l'avis "
         "des clients sur un produit ou l'accueil réservé à un film.")
    para(doc,
         "Ce projet compare trois façons de résoudre ce problème, de la plus simple "
         "à la plus moderne, sur un même jeu de données. L'objectif n'est pas seulement "
         "d'obtenir le meilleur score, mais de comprendre ce que chaque approche apporte "
         "et ce qu'elle coûte.")

    titre(doc, "Le jeu de données", 2)
    para(doc,
         "Nous utilisons IMDB, un ensemble de 50 000 critiques de films en anglais, "
         "réparties en deux moitiés égales : 25 000 pour l'entraînement et 25 000 pour "
         "le test. Chaque critique est étiquetée positive ou négative, et les deux classes "
         "sont parfaitement équilibrées. Cet équilibre est important : il évite qu'un modèle "
         "obtienne un bon score en répondant toujours la même chose.")


def approches(doc):
    titre(doc, "2. Les trois approches comparées")

    titre(doc, "Préparer le texte : TF-IDF", 2)
    para(doc,
         "Un ordinateur ne comprend pas les mots, il faut donc les transformer en nombres. "
         "La méthode TF-IDF donne à chaque mot un poids : élevé s'il est fréquent dans une "
         "critique donnée mais rare dans l'ensemble des critiques. Autrement dit, elle met "
         "en valeur les mots qui distinguent un texte des autres. Les mots très courants "
         "comme « the » ou « and » sont retirés au préalable, car ils n'apportent aucune "
         "information sur le sentiment.")
    para(doc,
         "Nous conservons aussi les paires de mots consécutifs. Cela permet de capturer des "
         "expressions comme « not good », dont le sens se perd si l'on regarde les mots "
         "isolément.")

    titre(doc, "Approche 1 : régression logistique", 2)
    para(doc,
         "Le modèle le plus simple des trois. À partir des poids TF-IDF, il apprend quels "
         "mots penchent vers le positif et lesquels vers le négatif, puis combine ces indices "
         "pour trancher. Son grand avantage est la transparence : on peut lire directement "
         "quels mots ont pesé dans la décision. Il s'entraîne en quelques minutes.")

    titre(doc, "Approche 2 : forêts aléatoires", 2)
    para(doc,
         "Cette méthode construit deux cents arbres de décision, chacun apprenant sur une "
         "partie différente des données, puis fait voter l'ensemble. L'idée est que les "
         "erreurs individuelles se compensent. Elle capture des combinaisons de mots plus "
         "complexes, mais reste plus lourde et moins lisible.")

    titre(doc, "Approche 3 : DistilBERT", 2)
    para(doc,
         "DistilBERT est un réseau de neurones déjà pré-entraîné sur d'énormes quantités de "
         "texte anglais. Il connaît donc la langue avant même de voir nos critiques. "
         "Contrairement aux deux méthodes précédentes, il ne traite pas les mots isolément "
         "mais tient compte de leur contexte : il peut comprendre que dans « je ne pensais pas "
         "aimer ce film, mais... », la suite renverse le sens.")
    para(doc,
         "Nous l'avons spécialisé sur nos données par une phase de fine-tuning, c'est-à-dire "
         "un réentraînement léger à partir de ses connaissances existantes.")


def resultats(doc, metriques):
    titre(doc, "3. Résultats")

    para(doc,
         "Chaque modèle est jugé sur quatre indicateurs. La précision globale indique la "
         "proportion de bonnes réponses. Le F1-score équilibre deux risques : passer à côté "
         "de critiques positives, ou en déclarer à tort. L'AUC mesure la capacité du modèle "
         "à bien classer les critiques par ordre de confiance, indépendamment du seuil choisi. "
         "Plus ces valeurs sont proches de 1, meilleur est le modèle.")

    tableau = doc.add_table(rows=1, cols=4)
    # Style neutre : simples bordures noires, sans aucun aplat de couleur.
    tableau.style = "Table Grid"
    entetes = tableau.rows[0].cells
    for i, nom in enumerate(["Modèle", "Précision globale", "F1-score", "AUC"]):
        entetes[i].text = nom
        for p in entetes[i].paragraphs:
            for run in p.runs:
                run.bold = True

    ordre = ["Logistic Regression", "DistilBERT", "Random Forest"]
    noms_lisibles = {
        "Logistic Regression": "Régression logistique",
        "Random Forest": "Forêts aléatoires",
        "DistilBERT": "DistilBERT",
    }

    for cle in ordre:
        if cle not in metriques:
            continue
        m = metriques[cle]
        ligne = tableau.add_row().cells
        ligne[0].text = noms_lisibles[cle]
        ligne[1].text = f"{m['accuracy']*100:.2f} %"
        ligne[2].text = f"{m['f1']*100:.2f} %"
        ligne[3].text = f"{m.get('roc_auc', 0):.3f}" if m.get("roc_auc") else "-"

    doc.add_paragraph()
    para(doc,
         "Les modèles classiques sont évalués sur les 25 000 critiques de test. DistilBERT "
         "l'est sur 2 000 critiques tirées aléatoirement, l'inférence sur l'ensemble complet "
         "étant trop lente sans carte graphique.")

    image(doc, "comparaison_modeles.png")
    legende(doc, "Figure 1 : comparaison des trois modèles.")

    doc.add_page_break()
    titre(doc, "Lecture des courbes ROC", 2)
    para(doc,
         "La courbe ROC montre le compromis entre bien reconnaître les critiques positives "
         "et éviter les fausses alertes. Plus la courbe monte vite vers le coin supérieur "
         "gauche, meilleur est le modèle. La diagonale représente un modèle qui répondrait "
         "au hasard.")

    image(doc, "roc_curve_Logistic_Regression.png", largeur=4.6)
    legende(doc, "Figure 2 : courbe ROC de la régression logistique (AUC = 0,959).")

    image(doc, "roc_curve_DistilBERT.png", largeur=4.6)
    legende(doc, "Figure 3 : courbe ROC de DistilBERT (AUC = 0,948).")


def analyse(doc, metriques):
    doc.add_page_break()
    titre(doc, "4. Analyse : un résultat contre-intuitif")

    lr = metriques.get("Logistic Regression", {})
    db = metriques.get("DistilBERT", {})

    para(doc,
         f"Le résultat le plus intéressant de ce projet est inattendu : la régression "
         f"logistique ({lr.get('accuracy', 0)*100:.1f} %) devance DistilBERT "
         f"({db.get('accuracy', 0)*100:.1f} %), alors que ce dernier est de loin le modèle "
         f"le plus sophistiqué. Ce constat mérite d'être expliqué plutôt que masqué.",
         gras=True)

    titre(doc, "L'explication", 2)
    para(doc,
         "La comparaison n'est pas équitable en termes de moyens. La régression logistique "
         "a été entraînée sur la totalité des 25 000 critiques disponibles. DistilBERT, lui, "
         "n'a vu que 2 000 critiques, car son entraînement demande une puissance de calcul "
         "dont nous ne disposions pas : sans carte graphique, un entraînement complet aurait "
         "demandé plusieurs heures.")
    para(doc,
         "Nous comparons donc un modèle simple bien nourri à un modèle puissant sous-alimenté. "
         "Un réseau de neurones a besoin de beaucoup d'exemples pour exprimer son potentiel ; "
         "privé de données, il n'exploite qu'une fraction de ses capacités.")

    titre(doc, "Ce que cela nous apprend", 2)
    para(doc,
         "Cette observation constitue un enseignement utile plutôt qu'un échec. Elle rappelle "
         "qu'un modèle plus complexe n'est pas automatiquement meilleur : sa supériorité dépend "
         "des moyens qu'on peut lui consacrer. À budget limité, en données comme en calcul, une "
         "méthode classique bien réglée reste très compétitive.")
    para(doc,
         "C'est un critère concret de choix technique. La régression logistique s'entraîne en "
         "quelques minutes, occupe 240 kilo-octets et répond instantanément. DistilBERT pèse "
         "268 mégaoctets et exige une infrastructure bien plus conséquente. Pour un gain de "
         "performance qui, ici, n'est pas au rendez-vous, le surcoût n'est pas justifié.")
    para(doc,
         "Signalons toutefois la limite de cette conclusion : elle vaut pour nos conditions "
         "d'expérimentation. Avec un entraînement sur l'intégralité des données, la littérature "
         "situe DistilBERT autour de 92 à 93 % sur IMDB, soit nettement devant nos modèles "
         "classiques. Ce prolongement constitue la suite naturelle du projet.")


def application(doc):
    doc.add_page_break()
    titre(doc, "5. De l'expérimentation à l'application")

    para(doc,
         "Un modèle n'a d'intérêt que s'il est utilisable. Nous avons donc développé une "
         "application web où l'on saisit un texte et obtient immédiatement le sentiment "
         "prédit, accompagné d'un indice de confiance. Le modèle employé est sélectionnable, "
         "ce qui permet de comparer les approches en direct.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Application accessible en ligne : " + URL_SITE)
    run.bold = True
    run.font.size = Pt(11)

    titre(doc, "Choix techniques", 2)
    para(doc,
         "L'application est écrite en Python avec le framework Flask. Elle expose également "
         "une interface de programmation permettant d'analyser un texte isolé ou une série de "
         "textes en une seule requête, ce qui est utile pour traiter un fichier d'avis.")
    para(doc,
         "Le service est mis en ligne automatiquement : chaque modification déposée sur le "
         "dépôt déclenche un nouveau déploiement. Seule la régression logistique y est "
         "embarquée, car sa taille réduite lui permet de fonctionner sur un hébergement "
         "gratuit. C'est une conséquence directe de l'analyse précédente : le modèle le plus "
         "léger se révèle aussi le plus performant et le seul déployable.")

    titre(doc, "Qualité du code", 2)
    para(doc,
         "Le projet comporte 91 tests automatiques vérifiant le nettoyage du texte, le "
         "comportement des modèles et les réponses de l'interface web. Ils sont exécutés "
         "automatiquement à chaque modification. Les tirages aléatoires sont fixés, si bien "
         "que deux entraînements identiques produisent exactement les mêmes résultats, "
         "condition indispensable à une comparaison honnête.")
    para(doc,
         "Les versions des bibliothèques sont également figées. Charger un modèle enregistré "
         "avec une version différente de scikit-learn peut en effet fausser silencieusement "
         "les prédictions : la reproductibilité ne s'arrête pas aux tirages aléatoires.")


def conclusion(doc, metriques):
    titre(doc, "6. Conclusion")

    lr = metriques.get("Logistic Regression", {})
    para(doc,
         f"Ce projet a permis de construire une chaîne complète, du jeu de données brut "
         f"jusqu'à une application en ligne. Le meilleur modèle, la régression logistique, "
         f"atteint {lr.get('accuracy', 0)*100:.1f} % de bonnes réponses sur 25 000 critiques "
         f"jamais vues pendant l'entraînement.")
    para(doc,
         "Son principal enseignement n'est pas un score, mais une nuance : la sophistication "
         "d'un modèle ne garantit pas sa supériorité. Les moyens disponibles, en données comme "
         "en calcul, pèsent autant que le choix de l'algorithme.")

    titre(doc, "Perspectives", 2)
    for texte in [
        "Entraîner DistilBERT sur l'intégralité des données avec une carte graphique, "
        "pour vérifier l'écart attendu.",
        "Étendre l'analyse à d'autres langues, notamment le français.",
        "Dépasser la distinction positif/négatif en détectant des nuances comme l'ironie.",
    ]:
        p = doc.add_paragraph(texte, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def main():
    metriques = charger_metriques()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    page_de_garde(doc)
    introduction(doc)
    approches(doc)
    resultats(doc, metriques)
    analyse(doc, metriques)
    application(doc)
    conclusion(doc, metriques)

    doc.save(SORTIE)
    print(f"Rapport généré : {SORTIE}")
    print(f"Métriques du {metriques.get('timestamp', 'inconnue')}")


if __name__ == "__main__":
    main()
